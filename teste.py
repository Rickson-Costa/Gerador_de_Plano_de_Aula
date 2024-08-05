import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx2pdf import convert
from PyPDF2 import PdfMerger

def converter_docs_para_pdfs(lista_de_documentos):
    pdfs = []
    
    # Converter cada documento Word para PDF
    for i, doc_path in enumerate(lista_de_documentos):
        pdf_path = f'documento_{i}.pdf'
        convert(doc_path, pdf_path)
        pdfs.append(pdf_path)
    
    return pdfs

def juntar_pdfs(pdfs, pdf_final_path):
    merger = PdfMerger()
    
    # Adicionar cada PDF ao arquivo final
    for pdf in pdfs:
        merger.append(pdf)
    
    merger.write(pdf_final_path)
    merger.close()

def aplicar_estilo_run(run):
    # Define a fonte para "Times New Roman" e tamanho 9
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

def preencher_documento(doc, inputs):
    # Substitui texto nos parágrafos
    for paragraph in doc.paragraphs:
        for key, value in inputs.items():
            pattern = r'\b' + re.escape(key) + r'\b'
            if re.search(pattern, paragraph.text):
                new_paragraph_text = re.sub(pattern, value, paragraph.text)
                paragraph.clear()  # Remove o texto original
                run = paragraph.add_run(new_paragraph_text)
                aplicar_estilo_run(run)

    # Substitui texto nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in inputs.items():
                        pattern = r'\b' + re.escape(key) + r'\b'
                        if re.search(pattern, paragraph.text):
                            new_paragraph_text = re.sub(pattern, value, paragraph.text)
                            paragraph.clear()  # Remove o texto original
                            run = paragraph.add_run(new_paragraph_text)
                            aplicar_estilo_run(run)

def inserir_imagens(doc, imagem1_path, imagem2_path):
    if doc.tables:
        table = doc.tables[0]
        
        if len(table.rows) > 0:
            row = table.rows[0]
            
            # Adiciona a primeira imagem (logo_principal) na primeira célula da primeira linha
            cell1 = row.cells[0]
            cell1_paragraph = cell1.paragraphs[0]
            run1 = cell1_paragraph.add_run()
            run1.add_picture(imagem1_path)
            cell1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adiciona a segunda imagem (logo_secundaria) na segunda célula da mesma linha com redimensionamento
            if len(row.cells) > 1:
                cell2 = row.cells[1]
                cell2_paragraph = cell2.paragraphs[0]
                run2 = cell2_paragraph.add_run()
                run2.add_picture(imagem2_path, width=Inches(1))  # Ajuste o tamanho conforme necessário
                cell2_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row.add_cell()  # Adiciona uma nova célula
                cell2 = row.cells[1]
                cell2_paragraph = cell2.paragraphs[0]
                run2 = cell2_paragraph.add_run()
                run2.add_picture(imagem2_path, width=Inches(1))  # Ajuste o tamanho conforme necessário
                cell2_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def gerar_documento(modelo_path, saida_path, imagem1_path, imagem2_path, i):
    try:
        doc = Document(modelo_path)
    except FileNotFoundError:
        print(f"Erro: O arquivo modelo '{modelo_path}' não foi encontrado.")
        return
    except PermissionError:
        print(f"Erro: Sem permissão para acessar o arquivo modelo '{modelo_path}'.")
        return
    
    print()
    response = input('Digite as informações geradas pelo chatGPT: ')
    
    # Regex para extrair data
    match = re.search(r'(\d{2})/(\d{2})/(\d{4})', response)

    if match:
        dia = match.group(1)
        mes = match.group(2)
        ano = match.group(3)
    else:
        print('Data não encontrada na resposta.')
    response_list = response.split('#')

    if len(response_list) != 12:
        print("Número incorreto de respostas. Certifique-se de fornecer todas as informações.")
        return

    inputs = {
        "acolhida_diaria": response_list[1].strip(),
        "leitura_deleite": response_list[2].strip(),
        "unid_tem_1": response_list[3].strip(),
        "obj_geral_1": response_list[4].strip(),
        "BNCC1": response_list[5].strip(),
        "unid_tem_2": response_list[6].strip(),
        "obj_geral_2": response_list[7].strip(),
        "BNCC2": response_list[8].strip(),
        "unid_tem_3": response_list[9].strip(),
        "obj_geral_3": response_list[10].strip(),
        "BNCC3": response_list[11].strip(),
        "xx": dia,
        "yy": mes,
        "20zz": ano,
    }
    inserir_imagens(doc, imagem1_path, imagem2_path)
    preencher_documento(doc, inputs)


    try:
        doc.save(saida_path)
        print(f"\nDocumento gerado com sucesso: {saida_path}")
        return f'documento_preenchido{i}.docx'
    
    except PermissionError:
        print(f"Erro: Sem permissão para salvar o arquivo em '{saida_path}'.")

if __name__ == "__main__":
    current_directory = os.path.dirname(os.path.abspath(__file__))
    modelo_path = os.path.join(current_directory, 'modelo.docx')
    
    imagem1_path = os.path.join(current_directory, 'logo_principal.png')
    imagem2_path = os.path.join(current_directory, 'logo_secundaria.png')
    status = int(input('Quantas páginas serão geradas? '))
    lista_de_documento = []
    for i in range(status):
        saida_path = os.path.join(current_directory, f'documento_preenchido{i}.docx')
        lista_de_documento.append(gerar_documento(modelo_path, saida_path, imagem1_path, imagem2_path, i))
    pdf_path = 'plano_de_aula.pdf'

    # Converter documentos Word para PDFs
    pdfs = converter_docs_para_pdfs(lista_de_documento)
    
    # Caminho do PDF final
    pdf_final_path = 'plano_de_aula.pdf'
    
    # Juntar todos os PDFs em um único arquivo
    juntar_pdfs(pdfs, pdf_final_path)
    
    # Apagar arquivos PDF temporários
    for pdf in pdfs:
        os.remove(pdf)
    

    # Apagar os documentos Word
    for doc in lista_de_documento:
        os.remove(doc)

    print(f'Arquivo PDF gerado: {pdf_path}')
